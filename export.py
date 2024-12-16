from sqlalchemy import func
from models import db, Details, Guide, Status, Department
from docx import *
from docx.shared import *
from docx.enum.text import *
from datetime import date, datetime
import pandas as pd

class ScholarDataHandler:
    def __init__(self):
        pass  # SQLAlchemy handles connection automatically

    def category_sort(self, department_name, gender):
        # Helper function to get the count based on conditions
        def get_count(religion, caste):
            query = db.session.query(func.count().label('count')).filter(
                Details.guide.has(Guide.department.has(dname=department_name)),
                Details.gender == gender,
                Details.religion == religion
            )
            if isinstance(caste, list):
                query = query.filter(Details.caste.in_(caste))
            else:
                query = query.filter(Details.caste == caste)
            return query.scalar()

        json_data = {
            'hindu': [
                get_count('hindu', 'oc'),
                get_count('hindu', 'sc'),
                get_count('hindu', 'st'),
                get_count('hindu', ['bc', 'mbc', 'dnc'])
            ],
            'christian': [
                get_count('christian', 'oc'),
                get_count('christian', 'sc'),
                get_count('christian', 'st'),
                get_count('christian', ['bc', 'mbc', 'dnc'])
            ],
            'muslim': [
                get_count('muslim', 'oc'),
                get_count('muslim', 'sc'),
                get_count('muslim', 'st'),
                get_count('muslim', ['bc', 'mbc', 'dnc'])
            ]
        }
        return json_data

class CatagoryExporter:
    def __init__(self, doc):
        self.doc = doc
        self.headings = ["S.NO", "Department", "Religion", "OC", "SC", "ST", "MBC/DNC/BC"]
        self.width_head = [0.54, 1.11, 0.85, 1.28,1.28,1.28,1.28,]
        self.database = ScholarDataHandler()

    
    def add_heading(self, Heading):
        heading = self.doc.add_paragraph()
        run = heading.add_run(Heading)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = "Book Antiqua"
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    def create_table(self):
        table = self.doc.add_table(rows=1, cols=7)
        table.style = 'Table Grid'
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = Inches(self.width_head[i])
                cell.text = self.headings[i]
                content = cell.paragraphs[0]
                content_run = content.runs[0]
                # ell Content Formatting
                content_run.font.size = Pt(12)
                content_run.font.bold = True
                content.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        return table
    
    def insert_data(self, table, departments_list, gender):
        for index, department in enumerate(departments_list):
            data = self.database.category_sort(department_name=department, gender=gender)
            self.row_1 = table.add_row()
            self.row_2 = table.add_row()
            self.row_3 = table.add_row()

            self.row_1.cells[0].merge(self.row_2.cells[0].merge(self.row_3.cells[0]))
            self.row_1.cells[1].merge(self.row_2.cells[1].merge(self.row_3.cells[1]))


            self.row_1.cells[0].text = str(index+1)
            self.row_1.cells[1].text = department
            self.row_1.cells[2].text = "Hindu"
            self.row_1.cells[3].text = str(data["hindu"][0])
            self.row_1.cells[4].text = str(data["hindu"][1])
            self.row_1.cells[5].text = str(data["hindu"][2])
            self.row_1.cells[6].text = str(data["hindu"][3])

            self.row_2.cells[2].text = "Christian"
            self.row_2.cells[3].text = str(data["christian"][0])
            self.row_2.cells[4].text = str(data["christian"][1])
            self.row_2.cells[5].text = str(data["christian"][2])
            self.row_2.cells[6].text = str(data["christian"][3])

            self.row_3.cells[2].text = "Muslim"
            self.row_3.cells[3].text = str(data["muslim"][0])
            self.row_3.cells[4].text = str(data["muslim"][1])
            self.row_3.cells[5].text = str(data["muslim"][2])
            self.row_3.cells[6].text = str(data["muslim"][3])


class AwardedExporter:
    def __init__(self, session, doc=None):
        self.session = session
        self.doc = doc or Document()
        self.headings = ["S.NO", "Year", "Gender", "Part-Time", "Full-Time", "Total Awarded"]
        self.width_head = [0.54, 1.11, 0.85, 1.23, 1.23, 1.8]

    def add_heading(self, heading):
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(heading)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = "Book Antiqua"
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def create_table(self):
        table = self.doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = Inches(self.width_head[i])
                cell.text = self.headings[i]
                content = cell.paragraphs[0]
                content_run = content.runs[0]
                content_run.font.size = Pt(12)
                content_run.font.bold = True
                content.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        return table

    def insert_data(self, data, table, year, index):
        row_1 = table.add_row()
        row_2 = table.add_row()
        
        row_1.cells[0].merge(row_2.cells[0])
        row_1.cells[1].merge(row_2.cells[1])
        row_1.cells[5].merge(row_2.cells[5])
        
        row_1.cells[0].text = str(index)
        row_1.cells[1].text = str(year)
        row_1.cells[2].text = "Male"
        row_1.cells[3].text = str(data["male"][0])
        row_1.cells[4].text = str(data["male"][1])
        row_1.cells[5].text = str(data["male"][0] + data["male"][1] + data["female"][0] + data["female"][1])

        row_2.cells[2].text = "Female"
        row_2.cells[3].text = str(data["female"][0])
        row_2.cells[4].text = str(data["female"][1])

    def get_awarded_year(self, department_name):
        year_list = (
            self.session.query(func.year(Details.viva_date))
            .select_from(Details)
            .join(Guide, Guide.gno == Details.gno)  # Join Details with Guide
            .join(Department, Department.dno == Guide.dno)  # Join Guide with Department
            .join(Status, Details.statusno == Status.statusno)  # Join Details with Status
            .filter(Details.viva_date.isnot(None))
            .filter(Status.current_status == "Awarded")  # Filter by Status
            .filter(Department.dname == department_name)  # Filter by Department name
            .distinct()
            .all()
        )
        
        # Return distinct years as a list, removing any None values
        return [year[0] for year in year_list if year[0] is not None]

    def awarded_sort(self, year, department_name):
        male_pt = (
            self.session.query(func.count())
            .select_from(Details)
            .join(Guide, Guide.gno == Details.gno)  # Join Details with Guide
            .join(Department, Department.dno == Guide.dno)  # Join Guide with Department
            .join(Status, Details.statusno == Status.statusno)  # Join Details with Status
            .filter(func.year(Details.viva_date) == year)
            .filter(Details.timing == "Part-Time")
            .filter(Details.gender == "Male")
            .filter(Department.dname == department_name)
            .filter(Status.current_status == "Awarded")  # Filter by Status
            .scalar()
        )
        male_ft = (
            self.session.query(func.count())
            .select_from(Details)
            .join(Guide, Guide.gno == Details.gno)  # Join Details with Guide
            .join(Department, Department.dno == Guide.dno)  # Join Guide with Department
            .join(Status, Details.statusno == Status.statusno)  # Join Details with Status
            .filter(func.year(Details.viva_date) == year)
            .filter(Details.timing == "Full-Time")
            .filter(Details.gender == "Male")
            .filter(Department.dname == department_name)
            .filter(Status.current_status == "Awarded")  # Filter by Status
            .scalar()
        )
        female_pt = (
            self.session.query(func.count())
            .select_from(Details)
            .join(Guide, Guide.gno == Details.gno)  # Join Details with Guide
            .join(Department, Department.dno == Guide.dno)  # Join Guide with Department
            .join(Status, Details.statusno == Status.statusno)  # Join Details with Status
            .filter(func.year(Details.viva_date) == year)
            .filter(Details.timing == "Part-Time")
            .filter(Details.gender == "Female")
            .filter(Department.dname == department_name)
            .filter(Status.current_status == "Awarded")  # Filter by Status
            .scalar()
        )
        female_ft = (
            self.session.query(func.count())
            .select_from(Details)
            .join(Guide, Guide.gno == Details.gno)  # Join Details with Guide
            .join(Department, Department.dno == Guide.dno)  # Join Guide with Department
            .join(Status, Details.statusno == Status.statusno)  # Join Details with Status
            .filter(func.year(Details.viva_date) == year)
            .filter(Details.timing == "Full-Time")
            .filter(Details.gender == "Female")
            .filter(Department.dname == department_name)
            .filter(Status.current_status == "Awarded")  # Filter by Status
            .scalar()
        )
        return {
            'male': [male_pt, male_ft],
            'female': [female_pt, female_ft]
        }

    def unknown_awarded_sort(self, department_name):
        male_pt = (
            self.session.query(func.count())
            .select_from(Details)
            .join(Guide, Guide.gno == Details.gno)  # Join Details with Guide
            .join(Department, Department.dno == Guide.dno)  # Join Guide with Department
            .join(Status)
            .filter(Details.viva_date.is_(None))
            .filter(Details.timing == "Part-Time")
            .filter(Details.gender == "Male")
            .filter(Department.dname == department_name)
            .filter(Status.current_status == "Awarded")
            .scalar()
        )
        male_ft = (
            self.session.query(func.count())
            .select_from(Details)
            .join(Guide, Guide.gno == Details.gno)  # Join Details with Guide
            .join(Department, Department.dno == Guide.dno)  # Join Guide with Department
            .join(Status)
            .filter(Details.viva_date.is_(None))
            .filter(Details.timing == "Full-Time")
            .filter(Details.gender == "Male")
            .filter(Department.dname == department_name)
            .filter(Status.current_status == "Awarded")
            .scalar()
        )
        female_pt = (
            self.session.query(func.count())
            .select_from(Details)
            .join(Guide, Guide.gno == Details.gno)  # Join Details with Guide
            .join(Department, Department.dno == Guide.dno)  # Join Guide with Department
            .join(Status)
            .filter(Details.viva_date.is_(None))
            .filter(Details.timing == "Part-Time")
            .filter(Details.gender == "Female")
            .filter(Department.dname == department_name)
            .filter(Status.current_status == "Awarded")
            .scalar()
        )
        female_ft = (
            self.session.query(func.count())
            .select_from(Details)
            .join(Guide, Guide.gno == Details.gno)  # Join Details with Guide
            .join(Department, Department.dno == Guide.dno)  # Join Guide with Department
            .join(Status)
            .filter(Details.viva_date.is_(None))
            .filter(Details.timing == "Full-Time")
            .filter(Details.gender == "Female")
            .filter(Department.dname == department_name)
            .filter(Status.current_status == "Awarded")
            .scalar()
        )
        return {
            'male': [male_pt, male_ft],
            'female': [female_pt, female_ft]
        }

class UnpaidExporter:
    def __init__(self):
        self.without_lab_headings = ["S.NO", "Name of the Guide", "Name of the Scholars", "Tuition Fees Paid", "Tuition Fees Unpaid"]
        self.without_lab_width_head = [0.5, 1.94, 2, 1.5,1.19]
        self.with_lab_headings = ["S.NO", "Name of the Guide", "Name of the Scholars", "Tuition Fees Paid", "Tuition Fees Unpaid", "Lab Fees Paid", "Lab Fees Unpaid"]
        self.with_lab_width_head = [0.5, 1.94, 2, 1.5,1.19, 1, 1]
        self.todayDate = date.today()
    
    def add_heading(self, Heading):
        heading = self.doc.add_paragraph()
        run = heading.add_run("\n"+Heading)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = "Book Antiqua"
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    def create_without_lab_table(self):
        table = self.doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = Inches(self.without_lab_width_head[i])
                cell.text = self.without_lab_headings[i]
                content = cell.paragraphs[0]
                content_run = content.runs[0]
                # ell Content Formatting
                content_run.font.size = Pt(12)
                content_run.font.bold = True
                content.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        return table
    
    def insert_without_lab_data(self, data):
        table = self.create_without_lab_table()
        data = self.find_unpaid_without_lab_scholars(data)
        for sno, row_data in enumerate(data):
            row = table.add_row()
            row.cells[0].text = str(sno+1)+ "."
            for index, item in enumerate(row_data):
                row.cells[index+1].text = str(item)
    
    def create_with_lab_table(self):
        table = self.doc.add_table(rows=1, cols=7)
        table.style = 'Table Grid'
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = Inches(self.with_lab_width_head[i])
                cell.text = self.with_lab_headings[i]
                content = cell.paragraphs[0]
                content_run = content.runs[0]
                # ell Content Formatting
                content_run.font.size = Pt(12)
                content_run.font.bold = True
                content.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        return table
    
    def insert_with_lab_data(self, data):
        table = self.create_with_lab_table()
        data = self.find_unpaid_with_lab_scholars(data)
        for sno, row_data in enumerate(data):
            row = table.add_row()
            row.cells[0].text = str(sno+1)+ "."
            for index, item in enumerate(row_data):
                row.cells[index+1].text = str(item)

    
    def find_unpaid_without_lab_scholars(self, details):
        unpaid_scholars = []
        for scholar in details:
            if int(scholar[4]) != 0:
                unpaid_scholars.append(scholar[0:4])
        return unpaid_scholars
    
    def find_unpaid_with_lab_scholars(self, details):
        unpaid_scholars = []
        for scholar in details:
            if scholar[6] != 0:
                unpaid_scholars.append(scholar[0:6])
        return unpaid_scholars
    
    def get_doing_scholars(self, department, lab):
        if not lab:
            query = (
                db.session.query(
                    Guide.guide_name,
                    Details.scholar_name,
                    Details.tution_fees_paid,
                    Details.tution_fees_unpaid,
                    Details.total_fees_unpaid
                )
                .join(Guide, Details.gno == Guide.gno)  # Join Details with Guide
                .join(Department, Guide.dno == Department.dno)  # Join Guide with Department
                .join(Status, Details.statusno == Status.statusno)  # Join Details with Status
                .filter(
                    Department.dname == str(department),
                    Status.current_status == 'doing'
                )
            )

            # Execute the query and fetch results
            return query.all()
        else:
            query = (
                db.session.query(
                    Guide.guide_name,
                    Details.scholar_name,
                    Details.tution_fees_paid,
                    Details.tution_fees_unpaid,
                    Details.labfees_paid,
                    Details.labfees_unpaid,
                    Details.total_fees_unpaid
                )
                .join(Guide, Details.gno == Guide.gno)  # Join Details with Guide
                .join(Department, Guide.dno == Department.dno)  # Join Guide with Department
                .join(Status, Details.statusno == Status.statusno)  # Join Details with Status
                .filter(
                    Department.dname == str(department),
                    Status.current_status == 'doing'
                )
            )

            # Execute the query and fetch results
            return query.all()


class ExcelDataExport:
    def __init__(self, db, filename):
        self.engine = db  # SQLAlchemy engine to connect to the database
        self.filename = filename  # Output filename for the Excel file

    def export(self):        
        # Query to get all departments
        departments = self.engine.query(Department).all()

        # Create an Excel writer
        with pd.ExcelWriter(self.filename, engine='openpyxl') as writer:
            for department in departments:
                # Query to get scholars for the current department
                scholars = self.engine.query(Details).join(Guide).filter(Guide.dno == department.dno).all()

                # Prepare the data for the DataFrame
                data = []
                for scholar in scholars:
                    data.append({
                        'Scholar ID': scholar.scholar_id,
                        'Guide Name': scholar.guide.guide_name if scholar.guide else None,
                        'Department': department.dname,
                        'Scholar Name': scholar.scholar_name,
                        'Gender': scholar.gender,
                        'Reg No': scholar.regno,
                        'DOB': scholar.dob,
                        'Timing': scholar.timing,
                        'Caste': scholar.caste,
                        'Subcaste': scholar.subcaste,
                        'Religion': scholar.religion,
                        'Email': scholar.email,
                        'Phone Number': scholar.phno,
                        'Optional Phone Number': scholar.optionalphno,
                        'Address': scholar.address,
                        'Commencement Date': scholar.commencement_date,
                        'Join Date': scholar.join_date,
                        'Annual Fee': scholar.annual_fee,
                        'Total Tuition Fees': scholar.total_tution_fees,
                        'Total Lab Fees': scholar.total_lab_fees,
                        'Total Center Fees': scholar.total_center_fees,
                        'Last Fee Date': scholar.last_fee_date,
                        'Tuition Fees Paid': scholar.tution_fees_paid,
                        'Lab Fees Paid' : scholar.labfees_paid,
                        'Tuition Fees Unpaid': scholar.tution_fees_unpaid,
                        'Lab Fees Unpaid': scholar.labfees_unpaid,
                        'Total Fees Unpaid': scholar.total_fees_unpaid,
                        'No Due Date': scholar.no_due_date,
                        'Viva Date': scholar.viva_date,
                        'Thesis Title': scholar.thesis_title,
                        'Status': scholar.status.current_status if scholar.status else None,
                        '1st DC Date': scholar.first_dc_meet,
                        '1st DC Fee': scholar.first_dc_fee,
                        '2nd DC Date': scholar.second_dc_meet,
                        '2nd DC Fee': scholar.second_dc_fee,
                        '3rd DC Date': scholar.third_dc_meet,
                        '3rd DC Fee': scholar.third_dc_fee
                    })

                # Convert to DataFrame
                df_new = pd.DataFrame(data)

                # Write the DataFrame to a new sheet in the Excel file
                df_new.to_excel(writer, sheet_name=department.dname[:31], index=False)  # Excel sheet names max 31 chars




# class UnpaidExporter:
#     def __init__(self):
#         self.without_lab_headings = ["S.NO", "Name of the Guide", "Name of the Scholars", "Tuition Fees Paid", "Tuition Fees Unpaid"]
#         self.without_lab_width_head = [0.5, 1.94, 2, 1.5,1.19]
#         self.with_lab_headings = ["S.NO", "Name of the Guide", "Name of the Scholars", "Tuition Fees Paid", "Tuition Fees Unpaid", "Lab Fees Paid", "Lab Fees Unpaid"]
#         self.with_lab_width_head = [0.5, 1.94, 2, 1.5,1.19, 1, 1]
#         self.todayDate = date.today()
    
#     def add_heading(self, Heading):
#         heading = self.doc.add_paragraph()
#         run = heading.add_run("\n"+Heading)
#         run.bold = True
#         run.font.size = Pt(12)
#         run.font.name = "Book Antiqua"
#         heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
#     def create_without_lab_table(self):
#         table = self.doc.add_table(rows=1, cols=5)
#         table.style = 'Table Grid'
#         for row in table.rows:
#             for i, cell in enumerate(row.cells):
#                 cell.width = Inches(self.without_lab_width_head[i])
#                 cell.text = self.without_lab_headings[i]
#                 content = cell.paragraphs[0]
#                 content_run = content.runs[0]
#                 # ell Content Formatting
#                 content_run.font.size = Pt(12)
#                 content_run.font.bold = True
#                 content.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#         return table
    
#     def insert_without_lab_data(self, data):
#         table = self.create_without_lab_table()
#         data = self.find_unpaid_without_lab_scholars(data)
#         for sno, row_data in enumerate(data):
#             row = table.add_row()
#             row.cells[0].text = str(sno+1)+ "."
#             for index, item in enumerate(row_data):
#                 row.cells[index+1].text = str(item)
    
#     def create_with_lab_table(self):
#         table = self.doc.add_table(rows=1, cols=7)
#         table.style = 'Table Grid'
#         for row in table.rows:
#             for i, cell in enumerate(row.cells):
#                 cell.width = Inches(self.with_lab_width_head[i])
#                 cell.text = self.with_lab_headings[i]
#                 content = cell.paragraphs[0]
#                 content_run = content.runs[0]
#                 # ell Content Formatting
#                 content_run.font.size = Pt(12)
#                 content_run.font.bold = True
#                 content.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#         return table
    
#     def insert_with_lab_data(self, data):
#         table = self.create_with_lab_table()
#         data = self.find_unpaid_with_lab_scholars(data)
#         for sno, row_data in enumerate(data):
#             row = table.add_row()
#             row.cells[0].text = str(sno+1)+ "."
#             for index, item in enumerate(row_data):
#                 row.cells[index+1].text = str(item)

#     def compare_dates(self, day, month):
#         """Return True if the given date is after today."""
#         return month > self.todayDate.month or (month == self.todayDate.month and day > self.todayDate.day)
    
#     def find_unpaid_without_lab_scholars(self, details):
#         unpaid_scholars = []
#         for scholar in details:
#             # Calculate the number of years from enrollment
#             enrollment_year = scholar[4] - 1
#             years_diff = self.todayDate.year - enrollment_year
#             if self.compare_dates(scholar[2], scholar[3]):
#                 years_diff -= 1

#             total_fees_to_be_paid = int(scholar[5]) * years_diff

#             if total_fees_to_be_paid != int(scholar[6]):
#                 unpaid_scholars.append([
#                     scholar[0],  # Guide name
#                     scholar[1],  # Scholar name
#                     scholar[6],  # Fees paid
#                     int(total_fees_to_be_paid - int(scholar[6]))  # Unpaid fees
#                 ])
#         return unpaid_scholars
    
#     def find_unpaid_with_lab_scholars(self, details):
#         unpaid_scholars = []
#         for scholar in details:
#             enrollment_year = scholar[4] - 1
#             years_diff = self.todayDate.year - enrollment_year
#             if self.compare_dates(scholar[2], scholar[3]):
#                 years_diff -= 1

#             total_fees_to_be_paid = int(scholar[5]) * years_diff

#             if scholar[4] < 2020:
#                 if self.compare_dates(scholar[2], scholar[3]):
#                     lab_years_diff = self.todayDate.year - (2020 - 1) - 1
#                 else:
#                     lab_years_diff = self.todayDate.year - (2020 - 1)
#             else:
#                 if self.compare_dates(scholar[2], scholar[3]):
#                     lab_years_diff = self.todayDate.year - (scholar[4] - 1) - 1  
#                 else:
#                     lab_years_diff = self.todayDate.year - (scholar[4]-1)
            
#             total_lab_fee = lab_years_diff * 2000

#             if total_fees_to_be_paid != int(scholar[6]) or scholar[7] != total_lab_fee:
#                 unpaid_scholars.append([
#                     scholar[0],  # Guide name
#                     scholar[1],  # Scholar name
#                     scholar[6],  # Fees paid
#                     int(total_fees_to_be_paid - int(scholar[6])),  # Unpaid fees
#                     scholar[7],  # Lab fees paid
#                     total_lab_fee - scholar[7]  # Unpaid lab fees
#                 ])
#         return unpaid_scholars
    
#     def get_doing_scholars(self, department, lab):
#         if not lab:
#             query = (
#                 db.session.query(
#                     Guide.guide_name,
#                     Details.scholar_name,
#                     func.day(Details.join_date).label('day'),
#                     func.month(Details.join_date).label('month'),
#                     func.year(Details.join_date).label('year'),
#                     Details.annual_fee,
#                     Details.tution_fees_paid
#                 )
#                 .join(Guide, Details.gno == Guide.gno)  # Join Details with Guide
#                 .join(Department, Guide.dno == Department.dno)  # Join Guide with Department
#                 .join(Status, Details.statusno == Status.statusno)  # Join Details with Status
#                 .filter(
#                     Department.dname == str(department),
#                     Status.current_status == 'doing'
#                 )
#             )

#             # Execute the query and fetch results
#             return query.all()
#         else:
#             query = (
#                 db.session.query(
#                     Guide.guide_name,
#                     Details.scholar_name,
#                     func.day(Details.join_date).label('day'),
#                     func.month(Details.join_date).label('month'),
#                     func.year(Details.join_date).label('year'),
#                     Details.annual_fee,
#                     Details.tution_fees_paid,
#                     Details.labfees_paid
#                 )
#                 .join(Guide, Details.gno == Guide.gno)  # Join Details with Guide
#                 .join(Department, Guide.dno == Department.dno)  # Join Guide with Department
#                 .join(Status, Details.statusno == Status.statusno)  # Join Details with Status
#                 .filter(
#                     Department.dname == str(department),
#                     Status.current_status == 'doing'
#                 )
#             )

#             # Execute the query and fetch results
#             return query.all()