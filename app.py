from flask import Flask, render_template, jsonify, send_file, redirect, url_for, flash
from flask_admin import Admin
from flask_admin.menu import MenuLink
from flask_admin.contrib.sqla import ModelView
from views import *
from export import *
from docx import *
from docx.shared import *
from docx.enum.text import *
import os, json

config = json.loads(open("config.json", "r").read())


app = Flask(__name__)
app.config["SQLALCHEMY_DATABASE_URI"] = f"mysql+pymysql://root:{config["password"]}@localhost/{config["dbname"]}"
app.config['FLASK_ADMIN_FLUID_LAYOUT'] = True
app.config['WORD_FILES'] = os.path.join(os.path.dirname(__file__), 'word')
db.init_app(app)
admin = Admin(app,template_mode='bootstrap4', name="Phd Scholar Archive", index_view=CustomAdminModelView())
app.secret_key = "23i7jkb8734rjh"


admin.add_view(DepartmentModelView(Department, db.session, category="Tables"))
admin.add_view(GuideModelView(db.session, category="Tables"))
admin.add_view(StatusModelView(Status, db.session, category="Tables"))
admin.add_view(DetailsModelView(db.session, category="Tables"))
admin.add_view(FeesUpdateView(name="Update Fees", category="Fees"))
admin.add_link(MenuLink(name="Refresh",url='/api/refresh'))

@app.route("/api/unpaid")
def unpaid_scholars():
    with_lab_departments = [str(dept) for dept in db.session.query(Department).filter(Department.labincluded == 1).all()]
    without_lab_departments = [str(dept) for dept in db.session.query(Department).filter(Department.labincluded == 0).all()]

    exporter = UnpaidExporter()
    exporter.doc = Document()

    for department in without_lab_departments:
        exporter.add_heading(department)
        details = exporter.get_doing_scholars(department, False)
        exporter.insert_without_lab_data(details)
    
    for department in with_lab_departments:
        exporter.add_heading(department)
        details = exporter.get_doing_scholars(department, True)
        exporter.insert_with_lab_data(details)

    file_path = os.path.join(app.config['WORD_FILES'], "fees_unpaid_details.docx")
    exporter.doc.save(file_path)
    return send_file(file_path, as_attachment=True)

@app.route("/api/category_sort")
def category_sort():
    department_list = [str(dept) for dept in Department.query.all()]
    document = Document()
    table_word_file = CatagoryExporter(document)
    
    # Male Data
    table_word_file.add_heading("Male")
    male_table = table_word_file.create_table()
    table_word_file.insert_data(male_table, department_list, "Male")
    
    # Female Data
    table_word_file.add_heading("\nFemale")
    female_table = table_word_file.create_table()
    table_word_file.insert_data(female_table, department_list, "Female")

    file_path = os.path.join(app.config['WORD_FILES'], "caste_category_details.docx")
    document.save(file_path)
    return send_file(os.path.join(app.config['WORD_FILES'], "caste_category_details.docx"), as_attachment=True)


@app.route("/api/awarded_sort")
def awarded_sort():
    department_list = [str(dept) for dept in Department.query.all()]
    exporter = AwardedExporter(db.session)
    document = Document()
    exporter.doc = document

    for department in department_list:
        years = exporter.get_awarded_year(department)
        exporter.add_heading(department)
        table = exporter.create_table()

        for index, year in enumerate(years):
            data = exporter.awarded_sort(year, department)
            exporter.insert_data(data, table, year, index + 1)
        unknown_data = exporter.unknown_awarded_sort(department)
        exporter.insert_data(unknown_data, table, 'Unknown', len(years) + 1)

    file_path = os.path.join(app.config['WORD_FILES'], "awarded_details.docx")
    document.save(file_path)
    return send_file(file_path, as_attachment=True)
    
@app.route("/api/export_all")
def exportall():
    exporter = ExcelDataExport(db.session, os.path.join(app.config['WORD_FILES'], "scholars_details.xlsx"))
    exporter.export()
    return send_file(os.path.join(app.config['WORD_FILES'], "scholars_details.xlsx"), as_attachment=True)


@app.route('/api/get_guides_opt', methods=['GET'])
def get_guides_by_department():
    department_id = request.args.get('department')
    guides = Guide.query.filter_by(dno=department_id).all()
    # Generate HTML options for the guide select field
    options = ''.join(f'<option value="{guide.gno}">{guide.guide_name}</option>' for guide in guides)
    print(options)
    return options

@app.route("/api/get_guides", methods=["GET"])
def get_guides():
    department_id = request.args.get("department")
    guides = Guide.query.filter_by(dno=department_id).all()
    guide_list = [{"id": guide.gno, "name": guide.guide_name} for guide in guides]
    return jsonify(guide_list)


@app.route("/api/get_labsfees", methods=["GET"])
def get_labfees():
    department_id = request.args.get("department")
    department = Department.query.filter_by(dno=department_id).first()
    isincluded = department.labincluded if department else False  # Default to False
    return jsonify({"lab": isincluded})

@app.route("/api/get_guides_scholar")
def get_scholar_name():
    department_id = request.args.get("department")
    guides = Guide.query.filter_by(dno=department_id).all()

    # Collect options
    options = []
    for guide in guides:
        scholars = Details.query.filter_by(gno=guide.gno).all()
        for scholar in scholars:
            option = f'<option value="{scholar.scholar_id}">{scholar.scholar_name} - (id: {scholar.scholar_id}, Guide: {guide.guide_name})</option>'
            options.append(option)

    # Join collected options into a single string
    result = ''.join(options)
    print(result)
    return result

@app.route("/api/refresh", methods=["GET"])
def refresh_details():
    try:
        # Fetch all scholars' IDs
        scholars = db.session.execute(text("""
            SELECT scholar_id, join_date, annual_fee FROM details
        """)).fetchall()

        for scholar_id, join_date, annual_fee in scholars:
            # Check if lab is included
            result = db.session.execute(text("""
                SELECT d.labincluded, d.labfees_year, d.annual_labfees
                FROM department d
                JOIN details dt ON dt.dno = d.dno
                WHERE dt.scholar_id = :scholar_id
            """), {'scholar_id': scholar_id}).fetchone()

            # Calculate fees
            is_labincluded = result and result[0] == 1
            labfees_year = result[1] if is_labincluded else 0
            annual_labfees = result[2] if is_labincluded else 0

            # Update fees
            db.session.execute(text("""
                UPDATE details
                SET total_tution_fees = total_tution_fees(:join_date, :annual_fee),
                    total_lab_fees = total_labfees(:join_date, :annual_labfees, :labfees_year)
                WHERE scholar_id = :scholar_id
            """), {
                'join_date': join_date,
                'annual_fee': annual_fee,
                'annual_labfees': annual_labfees,
                'labfees_year': labfees_year,
                'scholar_id': scholar_id
            })

            # Refresh updated values
            updated_values = db.session.execute(text("""
                SELECT total_tution_fees, total_lab_fees, tution_fees_paid, labfees_paid
                FROM details 
                WHERE scholar_id = :scholar_id
            """), {'scholar_id': scholar_id}).fetchone()

            if updated_values:
                tutionfees, labfees, tutionfeespaid, labfeespaid = updated_values

                # Calculate unpaid fees
                total_center_fees = tutionfees + (labfees if is_labincluded else 0)
                tutionfees_unpaid = tutionfees - tutionfeespaid
                labfees_unpaid = labfees - labfeespaid if is_labincluded else 0
                total_fees_unpaid = tutionfees_unpaid + labfees_unpaid

                # Update unpaid fees
                db.session.execute(text("""
                    UPDATE details
                    SET total_center_fees = :total_center_fees,
                        tution_fees_unpaid = :tutionfees_unpaid,
                        labfees_unpaid = :labfees_unpaid,
                        total_fees_unpaid = :total_fees_unpaid
                    WHERE scholar_id = :scholar_id
                """), {
                    'total_center_fees': total_center_fees,
                    'tutionfees_unpaid': tutionfees_unpaid,
                    'labfees_unpaid': labfees_unpaid,
                    'total_fees_unpaid': total_fees_unpaid,
                    'scholar_id': scholar_id
                })

        db.session.commit()
        flash("Fees updation successful.")
        return redirect("/admin/details/")

    except Exception as e:
        db.session.rollback()
        return jsonify({"status": "error", "message": str(e)}), 500


if __name__ == "__main__":
    app.run(debug=True)

